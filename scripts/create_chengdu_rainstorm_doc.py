from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def set_run_font(run, size: float | None = None, bold: bool = False) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    run.bold = bold


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        set_run_font(run, 14 if level == 1 else 12, True)


def add_para(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, 10.5)


def main() -> None:
    out_dir = Path("test_documents")
    out_dir.mkdir(exist_ok=True)
    out = out_dir / "成都暴雨灾害深度研究测试文档.docx"

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("成都暴雨灾害深度研究测试文档")
    set_run_font(title_run, 20, True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("用于 SkyGuard Agent 深度研究 / GraphRAG / 风险评估流程测试")
    set_run_font(subtitle_run, 11)

    add_heading(document, "一、事件概况")
    add_para(
        document,
        "2026年7月上旬，成都市出现一次持续性强降雨过程，降雨主要影响中心城区、龙泉山西侧、天府新区、"
        "双流区、郫都区、新都区以及都江堰—彭州—崇州一带山前区域。本文档为测试材料，部分数值为模拟值，"
        "用于检验系统在灾害分析场景下的意图识别、联网检索、GraphRAG构建、人口暴露估算和报告生成能力。",
    )
    add_para(
        document,
        "本次过程具有短时雨强大、局地累积雨量高、城市内涝风险与山洪地质灾害风险并存的特点。"
        "建议 Agent 在分析时主动区分“已确认事实”“待核查信息”和“模型推断”。",
    )

    add_heading(document, "二、关键观测与模拟数据")
    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["区域", "累计雨量", "最大小时雨强", "主要风险", "备注"]
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = text
        shade(cell, "DDEBE7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    rows = [
        ["锦江区—青羊区—武侯区", "80—130 mm", "45 mm/h", "道路积水、地下空间倒灌", "人口密集，交通影响敏感"],
        ["双流区—天府新区", "100—160 mm", "55 mm/h", "低洼片区内涝、机场周边交通延误", "需关注机场高速、成雅高速"],
        ["都江堰—彭州—崇州", "120—210 mm", "70 mm/h", "山洪、滑坡、泥石流", "山前地带风险较高"],
        ["龙泉驿区", "90—150 mm", "50 mm/h", "城市径流、坡面汇流", "龙泉山附近需重点巡查"],
        ["新都区—郫都区", "70—120 mm", "40 mm/h", "农田渍涝、道路积水", "需关注排水泵站能力"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].text = text
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    add_heading(document, "三、已发现影响")
    impacts = [
        "部分主干道出现短时积水，低洼路段通行效率下降，公交线路可能出现临时绕行。",
        "地下车库、下穿隧道、地铁出入口周边存在雨水倒灌风险，需要重点排查排水设施。",
        "山前区域在强降雨叠加土壤含水量升高后，存在小型滑坡、崩塌和泥石流风险。",
        "河道、水库、排洪沟渠水位可能快速上涨，应关注锦江、府河、沙河及其支流沿线风险点。",
        "人口密集社区、学校、医院、养老机构和大型商业综合体需要纳入重点影响对象。",
    ]
    for item in impacts:
        document.add_paragraph(item, style="List Bullet")

    add_heading(document, "四、需要 Agent 主动检索核查的问题")
    questions = [
        "成都市气象台是否发布暴雨黄色、橙色或红色预警？预警发布时间和覆盖区域是什么？",
        "成都市应急管理、交通运输、水务或住建部门是否发布积水点、交通管制或避险转移信息？",
        "最近24小时真实雷达回波、降雨站点数据和未来6小时临近预报如何？",
        "本次降雨与历史典型成都暴雨过程相比处于什么等级？",
        "是否存在需要截图保存的权威网页、地图或公告页面？",
    ]
    for item in questions:
        document.add_paragraph(item, style="List Number")

    add_heading(document, "五、风险评估提示")
    add_para(
        document,
        "建议将风险分为城市内涝风险、山洪地质灾害风险、交通运行风险、重点人群暴露风险和次生灾害风险五类。"
        "若系统可读取本地人口密度缓存，应优先估算灾害点周边3公里、5公里、10公里范围内的暴露人口，"
        "并结合风险等级折算潜在受影响人口。",
    )
    add_para(
        document,
        "对于成都场景，中心城区人口密度较高，短时积水即可能造成较大社会影响；西部山地及山前地带虽然人口密度相对较低，"
        "但滑坡、泥石流等突发性风险更高，应在报告中分别表达。",
    )

    add_heading(document, "六、建议处置措施")
    suggestions = [
        "加强下穿隧道、地铁出入口、地下车库、低洼院落和老旧小区排水巡查。",
        "对山洪沟、地质灾害隐患点、临坡临崖道路和旅游景区开展临时管控。",
        "通过短信、政务新媒体、社区网格和交通诱导屏发布避险与绕行提示。",
        "对医院、学校、养老机构、避难场所和应急物资仓库进行重点保障。",
        "保留权威预警、降雨实况、交通管制和灾情核查信息作为报告证据。",
    ]
    for item in suggestions:
        document.add_paragraph(item, style="List Bullet")

    add_heading(document, "七、测试期望")
    add_para(
        document,
        "上传本文档后，深度研究工具应能够识别这是“灾害分析/深度研究”场景，而不是普通问答；"
        "应调用文档读取、GraphRAG/知识抽取、联网搜索、必要网页截图、风险评估和报告草稿生成相关能力。"
        "若需要生成正式报告，应先询问用户输出 Word 还是 PDF，并展示阶段性进度。",
    )
    add_para(
        document,
        "测试重点：是否能从文档抽取地点、灾害类型、影响对象、风险因素、待核查问题；是否能把对话记录与文档证据合并；"
        "是否能避免把模拟数据当作权威事实。",
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("SkyGuard 测试文档｜模拟材料，仅用于系统功能验证")
    set_run_font(footer_run, 9)

    document.save(out)
    print(out.resolve())


if __name__ == "__main__":
    main()
